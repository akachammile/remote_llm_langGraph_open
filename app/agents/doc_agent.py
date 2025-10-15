import traceback,json,re
from docx import Document
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from typing import Dict
from docx.oxml.ns import qn
from app.logger import logger
from docx.shared import Pt, Inches
from app.agents.base import BaseAgent
from app.graphs.graph_state import AgentState
from app.prompts.doc_prompt import SYSTEM_PROMPT
from langgraph.graph import StateGraph, START, END


class DocAgent(BaseAgent):
    name: str = "DocAgent"
    description: str = "用于将理文档相关的具体任务，可对文档进行读取, 也可生成DOCX文件"
    system_prompt: str = SYSTEM_PROMPT
    tool: Dict[str, str] = {
        "function_name": "analysis",
    }
    def __init__(self, output_path="分析结果.docx", ):
        super().__init__()
        self.output_path = r"E:\1_LLM_PROJECT\remote_llm_langGraph_open\temp\分析结果.docx"
        self.doc_template = r"E:\1_LLM_PROJECT\remote_llm_langGraph_open\app\template\xinjiang.docx"
        # self._load_template()
        # if not self.doc_template:
        #     self.doc_template = self._load_template()
    
    
    async def doc_agent_supervisor(self, state: AgentState)-> AgentState:
        parsed = await self._analyze_content(state)
        state["image_content"] = parsed
        state["next_agent"] = "doc"
        return state
        
    
    async def _analyze_content(self, state: AgentState) -> AgentState:
        """调用 LLM, 返回 JSON"""
        # FIXME 这里需要重新分析，这里如果没有图像，则需要大模型重新反思决策
        title_list = []
        document = Document(self.doc_template)
        for para in document.paragraphs:
            if para.style.name.startswith('标题'):
              title_list.append(para.text)
        
        prompt_template: str = self.system_prompt.\
            format(question = state["question"]
        , content = state.get("image_content", "无")
        , doc_formatt = title_list)

        response = await self.llm.ask_tool(prompt_template, state)
        json_str = response.content
        match = re.search(r"\{[\s\S]*\}", json_str)
        if match:
            json_str = match.group(0)
        return json.loads(json_str)
    def insert_paragraph_after(self, paragraph, text):
      """
      在指定的段落后面插入一个新段落。
      此版本修正了 OxmlElement 的兼容性问题。
      """
      # --- 修改之处 ---
      # 移除了旧版本不支持的 nsmap 参数
      new_p = OxmlElement("w:p") 
      # --- 修改结束 ---
      
      paragraph._p.addnext(new_p)
      new_run = new_p.makeelement(qn('w:r'))
      new_p.append(new_run)
      new_t = new_run.makeelement(qn('w:t'))
      new_run.append(new_t)
      new_t.text = text
    # def fill_content_by_headings(self, doc, content_map):
    #   """
    #   根据标题文本查找标题，并在其后填充内容。
    #   content_map 是一个字典 {heading_text: content_to_insert}
    #   """
    #   for para in doc.paragraphs:
    #       if para.style.name.startswith('标题') and para.text in content_map:
    #           print(f"找到标题: '{para.text}'，准备在其后填充内容...")
    #           content_to_add = content_map[para.text]
    #           self.insert_paragraph_after(para, content_to_add)

    
    async def write_doc(self, state: AgentState, title: str = "分析结果")->AgentState:
        content_to_fill = state["image_content"]
        document = Document(self.doc_template)
        cleaned_key_map = {re.sub(r'^\d+(\.\d+)*\s*', '', k).strip(): k for k in content_to_fill.keys()}

        for para in document.paragraphs:
            # 只处理标题段落
            if not para.style.name.startswith('标题'):
                continue
            
            para_text_cleaned = para.text.strip()

            # --- 调试输出 (非常重要) ---
            logger.debug(f"正在检查 DOCX 段落: '{para_text_cleaned}' (样式: {para.style.name})")
            
            # 使用清理过的文本进行匹配
            if para_text_cleaned in cleaned_key_map:
                # 找到匹配后，使用原始的、带编号的键来获取内容
                original_key = cleaned_key_map[para_text_cleaned]
                content_to_add = content_to_fill[original_key]
                
                logger.info(f"✅ 找到匹配: '{para_text_cleaned}' -> 准备填充内容...")
                
                # --- 填充内容的逻辑 ---
                new_p = OxmlElement("w:p") 
                para._p.addnext(new_p)
                new_run = new_p.makeelement(qn('w:r'))
                new_p.append(new_run)
                new_t = new_run.makeelement(qn('w:t'))
                new_run.append(new_t)
                new_t.text = content_to_add
            else:
                logger.warning(f"❌ 未找到匹配: '{para_text_cleaned}' 在 {list(cleaned_key_map.keys())} 中不存在。")
        document.save(self.output_path)
        state["next_agent"] = "exit"
        state["sub_task"] = "文档生成并写入任务已完成"
        state["processed_doc_path"] = [self.output_path]
        return state

    def build_subgraph(self):
        # 创建子图
        doc_subgraph = StateGraph(AgentState)
        doc_subgraph.add_node("analysis", self.doc_agent_supervisor)
        doc_subgraph.add_node("doc", self.write_doc)
        doc_subgraph.add_edge(START, "analysis")
        doc_subgraph.add_edge("analysis", "doc")
        doc_subgraph.add_edge("doc", END) 
        chat_subgraph = doc_subgraph.compile()
        return chat_subgraph